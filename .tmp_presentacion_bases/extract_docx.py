from docx import Document
from pathlib import Path
import sys

sys.stdout.reconfigure(encoding="utf-8")

source = Path(r"C:\Users\Carlos Barrientos\Downloads\MC-BC-IA Bases Convocatoria Reto IA PROSUR - Categorias por Áreas.docx")
doc = Document(source)

print("=== PARAGRAPHS ===")
for i, paragraph in enumerate(doc.paragraphs, 1):
    text = paragraph.text.strip()
    if text:
        print(f"P{i:03d} [{paragraph.style.name}]: {text}")

print("\n=== TABLES ===")
for ti, table in enumerate(doc.tables, 1):
    print(f"\nTABLE {ti} ({len(table.rows)}x{len(table.columns)})")
    for ri, row in enumerate(table.rows, 1):
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        print(f"R{ri:02d}: " + " || ".join(cells))

print("\n=== SECTIONS ===")
for si, section in enumerate(doc.sections, 1):
    print(f"SECTION {si}: {section.page_width} x {section.page_height}; margins {section.top_margin},{section.right_margin},{section.bottom_margin},{section.left_margin}")

print("\n=== SHAPES ===")
print(f"inline_shapes={len(doc.inline_shapes)}")
